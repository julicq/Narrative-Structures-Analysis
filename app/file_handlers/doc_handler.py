# app/file_handlers/doc_handler.py

import platform
import subprocess
import os
import logging
from docx import Document
from pathlib import Path
from typing import Optional
import docx2txt

logger = logging.getLogger(__name__)

def extract_doc_text(file_path: str) -> Optional[str]:
    """Извлечение текста из .doc файла с несколькими методами и fallback опциями
    
    Args:
        file_path (str): Путь к .doc файлу
        
    Returns:
        Optional[str]: Извлеченный текст или None в случае ошибки
        
    Raises:
        FileNotFoundError: Если файл не найден
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
        
    errors = []
    
    # MacOS: пробуем textutil
    if platform.system() == 'Darwin':
        try:
            txt_path = f"{file_path}.txt"
            subprocess.run(
                ['textutil', '-convert', 'txt', '-output', txt_path, file_path], 
                check=True,
                capture_output=True
            )
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
            os.remove(txt_path)
            return text
        except Exception as e:
            errors.append(f"textutil failed: {str(e)}")
            logger.warning(f"textutil conversion failed: {str(e)}, trying alternatives...")
    
    # Пробуем LibreOffice (для всех платформ)
    try:
        temp_dir = os.path.dirname(file_path)
        subprocess.run([
            'soffice',
            '--headless',
            '--convert-to', 'txt:Text',
            '--outdir', temp_dir,
            file_path
        ], check=True, capture_output=True)
        
        txt_file = os.path.join(temp_dir, f"{Path(file_path).stem}.txt")
        
        if os.path.exists(txt_file):
            with open(txt_file, 'r', encoding='utf-8') as f:
                text = f.read()
            os.remove(txt_file)
            return text
    except Exception as e:
        errors.append(f"LibreOffice failed: {str(e)}")
        logger.warning(f"LibreOffice conversion failed: {str(e)}, trying alternatives...")
    
    # Пробуем antiword
    try:
        result = subprocess.run(
            ['antiword', file_path], 
            capture_output=True, 
            text=True,
            check=True
        )
        if result.stdout.strip():
            return result.stdout
    except Exception as e:
        errors.append(f"antiword failed: {str(e)}")
        logger.warning(f"antiword conversion failed: {str(e)}, trying alternatives...")
    
    # Пробуем catdoc
    try:
        result = subprocess.run(
            ['catdoc', file_path], 
            capture_output=True, 
            text=True,
            check=True
        )
        if result.stdout.strip():
            return result.stdout
    except Exception as e:
        errors.append(f"catdoc failed: {str(e)}")
        logger.warning(f"catdoc conversion failed: {str(e)}")
    
    # Если все методы не сработали, логируем все ошибки и возвращаем None
    error_msg = "All conversion methods failed:\n" + "\n".join(errors)
    logger.error(error_msg)
    return None

def extract_docx_text(file_path: str) -> Optional[str]:
    """Извлечение текста из .docx файла с несколькими методами
    
    Args:
        file_path (str): Путь к .docx файлу
        
    Returns:
        Optional[str]: Извлеченный текст или None в случае ошибки
        
    Raises:
        FileNotFoundError: Если файл не найден
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Сначала пробуем docx2txt как более простой метод
    try:
        text = docx2txt.process(file_path)
        if text and text.strip():
            return text
    except Exception as e:
        logger.warning(f"docx2txt extraction failed: {str(e)}, trying python-docx...")
    
    # Если docx2txt не сработал, используем python-docx
    try:
        doc = Document(file_path)
        full_text = []
        
        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        # Проверяем, что мы действительно получили какой-то текст
        if full_text:
            return '\n'.join(full_text)
        
        logger.warning(f"No text content found in {file_path}")
        return None
        
    except Exception as e:
        logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
        return None

def extract_text(file_path: str) -> Optional[str]:
    """Универсальная функция для извлечения текста из .doc или .docx файлов
    
    Args:
        file_path (str): Путь к файлу
        
    Returns:
        Optional[str]: Извлеченный текст или None в случае ошибки
        
    Raises:
        ValueError: Если формат файла не поддерживается
        FileNotFoundError: Если файл не найден
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = file_path.suffix.lower()
    
    if extension == '.doc':
        return extract_doc_text(str(file_path))
    elif extension == '.docx':
        return extract_docx_text(str(file_path))
    else:
        error_msg = f"Unsupported file format: {extension}"
        logger.error(error_msg)
        raise ValueError(error_msg)
